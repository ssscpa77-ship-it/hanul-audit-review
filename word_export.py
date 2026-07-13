"""Word 문서 생성 — QC 자문의견·4대 중점 체크리스트 양식."""

from __future__ import annotations

from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Any

import fss_focus
import review_guidelines as rg

_FONT = "맑은 고딕"


def _doc():
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    return Document, WD_ALIGN_PARAGRAPH, Pt, RGBColor


def _heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _para(doc, text: str, *, bold: bool = False, size: int = 11) -> None:
    _, _, Pt, _ = _doc()
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = _FONT
    run.font.size = Pt(size)


def _bullet(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")


def _table(doc, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    for j, h in enumerate(headers):
        tbl.rows[0].cells[j].text = h
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            tbl.rows[i].cells[j].text = val


def build_advisory_opinion_docx() -> bytes:
    """인간 수준 리뷰노트 달성을 위한 QC·AI 설계 자문의견."""
    Document, WD_ALIGN_PARAGRAPH, Pt, RGBColor = _doc()
    doc = Document()
    today = date.today().strftime("%Y-%m-%d")

    title = doc.add_heading("감사조서 자가검토 — QC·AI 설계 자문의견", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para(doc, f"한울 회계법인 · hanul-002 · 작성일 {today}", size=10)
    _para(doc, "목적: 인간 심리(2차 검토) 회계사와 거의 동일한 수준의 리뷰노트 생성을 위한 설계·지침 보완 방향", bold=True)
    doc.add_paragraph()

    _heading(doc, "1. 핵심 결론", 1)
    _para(
        doc,
        "인간 회계사와 유사한 리뷰노트는 「키워드 검색기 + LLM」이 아니라, "
        "「증거가 연결된 조서 그래프 위에서 Tier1·2를 규칙으로 판정하고, "
        "AI는 부족한 문장만 다듬는 구조」로 달성할 수 있습니다.",
        bold=True,
    )
    _bullet(
        doc,
        [
            "L1 결정론: 절차·대사·4대 중점 트리거 (규칙)",
            "L2 증거 그래프: Lead↔세부↔조회서↔주석 연결 (해소 판정)",
            "L3 AI 서술: 미충족 항목의 reason·to_be만 생성",
            "학습 루프: 골든 조서 10~20건으로 오탐/미탐 수치 관리",
        ],
    )

    _heading(doc, "2. 인간 심리 리뷰 판정 순서", 1)
    _table(
        doc,
        ["단계", "회계사 질문", "현재 프로그램", "격차"],
        [
            ["A", "Engagement·중요성·리스크", "extract_engagement", "업종·특이리스크 미반영"],
            ["B", "4대 중점 해당?", "fss_focus", "키워드 체크 → 증거 패턴 필요"],
            ["C", "필수 절차 완료?", "PROCEDURE_RULES+KB", "키워드 1개=통과"],
            ["D", "조서 간 대사?", "교차합·Lead대사", "조서 역할 사전 확장 필요"],
            ["E", "판단·결론 충분?", "AI + qc_review", "그래프 기반 해소 판정"],
            ["F", "감리 유사 패턴?", "RAG 첨부", "역매칭 독립 지적"],
        ],
    )

    _heading(doc, "3. 규칙 보완 지침", 1)
    _heading(doc, "3.1 절차 규칙 — 키워드 → 증거 유형 체크리스트", 2)
    _para(doc, "절차 1건 정의 템플릿 (Hanul DB `계정별_필수절차_카탈로그`에 작성):")
    _table(
        doc,
        ["필드", "설명", "예(차입금)"],
        [
            ["procedure_id", "절차 식별자", "BB-01 외부조회"],
            ["trigger", "실행 조건", "차입금 > PM"],
            ["required_evidence", "필수 증거", "조회서·조회일·대사(True)"],
            ["acceptable_conclusion", "Pass 문장", "조회 결과 일치"],
            ["cross_sheet", "연결 조서", "BB100↔Lead↔CL"],
            ["fail_if", "무조건 Fail", "Ref만 있고 금액·결론 없음"],
        ],
    )

    _heading(doc, "3.2 교차합 — 비교 단위 정의", 2)
    _table(
        doc,
        ["계정군", "비교 단위", "제외"],
        [
            ["차입금", "유형 합계·합계행", "BB200 이자테스트, 5% 미만 소액"],
            ["우발부채", "조회서↔주석(제공/수혜 분리)", "상호 금액 직접 대사"],
            ["재고", "실사↔장부↔평가", "품목 vs 총계 직접 대사"],
        ],
    )

    _heading(doc, "3.3 검토내역 해소 판정", 2)
    _bullet(
        doc,
        [
            "해당 행·열에 검토·판단·결론 서술 + 업무상 타당한 결론 → 지적 제외",
            "단, 중대 절차 자체 누락은 문서화와 무관하게 지적 (우선순위 1)",
        ],
    )

    _heading(doc, "4. 4대 중점사항 — 이슈별 4단계 절차", 1)
    for step, q in rg.FOCUS_WORKFLOW_STEPS:
        _para(doc, f"{step}: {q}")
    _para(doc, "각 ChecklistItem에 required_evidence, acceptable_phrases, red_flag_phrases, ai_review_questions를 정의하면 항목별 Pass/Fail이 가능합니다.")

    _heading(doc, "5. AI 심리 리뷰노트 지침", 1)
    _para(doc, rg.REVIEW_NOTE_STYLE)
    doc.add_paragraph()
    _para(doc, rg.AI_ROLE_GUIDE)
    _bullet(
        doc,
        [
            "AI 호출: L1 Fail·4대 미충족·판단 계정만 (전수 호출 지양)",
            "금지: 조서에 없는 사실 단정, 근거 없는 기준번호, Lead에 있는 공시 재지적",
        ],
    )

    _heading(doc, "6. QC 3축 Tier 구조", 1)
    _table(
        doc,
        ["Tier", "성격", "카테고리"],
        [
            ["1", "4대 중점", "중점감리 · 상 · focus_protected"],
            ["2", "핵심 절차 누락", "절차누락 · 상"],
            ["3", "감리사례 보강", "기존 노트에 ⚖️ (향후 역매칭)"],
        ],
    )

    _heading(doc, "7. 실행 로드맵 (효율 순)", 1)
    _table(
        doc,
        ["Phase", "기간", "내용"],
        [
            ["A", "4~6주", "지침 3종 표 작성 → PROCEDURE_RULES·fss_focus 반영"],
            ["B", "6~10주", "증거 그래프 레이어 (조서 연결 통합)"],
            ["C", "지속", "골든셋 10~20건 회귀 (Tier1 recall≥95%)"],
            ["D", "B 이후", "AI 서술 전용 (호출량 1/3↓)"],
        ],
    )

    _heading(doc, "8. Hanul DB 템플릿 업로드 요청", 1)
    _para(doc, f"폴더: Hanul DB / {rg.GUIDELINES_DB_SUBDIR}/")
    rows = [[r["file"], r["purpose"], r["status"]] for r in rg.template_upload_requests()]
    _table(doc, ["파일명", "용도", "상태"], rows)

    _heading(doc, "9. 작업지시서 연동", 1)
    _bullet(
        doc,
        [
            "§A 인간 심리 리뷰 판정 트리",
            "§B 계정별 필수 절차 카탈로그",
            "§C 4대 중점 이슈별 검토 워크플로",
            "§D 조서 연결·대사 사전",
            "§E 검토내역·결론 인정 문장 사전",
            "§F 리뷰노트 문체·중요도",
            "§G AI 역할·호출 조건",
            "§H 골든셋·회귀 기준",
        ],
    )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _issue_template_section(
    doc,
    issue_no: int,
    title: str,
    related: str,
    checklist_names: list[str],
    listed: bool,
) -> None:
    src = "금융감독원(상장)" if listed else "한국공인회계사회(비상장)"
    _heading(doc, f"이슈 {issue_no}. {title}", 2)
    _para(doc, f"출처: {src} 4대 중점 · 관련 계정: {related or '(작성)'}")

    _para(doc, "■ 4단계 검토 워크플로 (작성란)", bold=True)
    _table(
        doc,
        ["단계", "검토 질문", "요구 증거 (작성)", "Pass 패턴 (작성)", "Fail 시 노트 유형"],
        [
            ["1 Trigger", "", "", "", "해당없음"],
            ["2 Procedure", "", "", "", "절차누락"],
            ["3 Focus", "", "", "", "중점감리"],
            ["4 Disclosure", "", "", "", "공시"],
        ],
    )
    doc.add_paragraph()

    _para(doc, "■ 체크리스트 항목별 상세 (작성란)", bold=True)
    headers = [
        "체크리스트항목",
        "관련조서",
        "필수키워드(detect_all)",
        "선택키워드(detect_any)",
        "필수증거",
        "인정문장(Pass)",
        "지적문장(Red flag)",
        "근거(basis)",
        "보완지침(to_be)",
        "AI검토질문",
    ]
    rows: list[list[str]] = []
    for name in checklist_names:
        rows.append([name] + [""] * (len(headers) - 1))
    # 빈 행 2개 추가
    for _ in range(2):
        rows.append([""] * len(headers))
    _table(doc, headers, rows)
    doc.add_paragraph()


def build_focus_checklist_templates_docx() -> bytes:
    """상장·비상장 4대 중점 이슈별 빈 양식표 Word."""
    Document, WD_ALIGN_PARAGRAPH, _, _ = _doc()
    doc = Document()
    h = doc.add_heading("4대 중점사항 체크리스트 작성 양식", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para(
        doc,
        "본 양식을 작성한 뒤 Hanul DB / 자가검토_지침_템플릿 / "
        "4대중점_체크리스트_상장(또는 비상장)_FY2026.xlsx 로 저장·업로드하십시오. "
        "프로그램이 자동 연동합니다.",
        size=10,
    )
    doc.add_page_break()

    _heading(doc, "【상장】 금융감독원 4대 중점", 1)
    for issue in fss_focus._DEFAULT_LISTED_2026:
        accts = ", ".join(issue.related_accounts) if issue.related_accounts else ""
        names = [c.name for c in issue.checklist] or [""]
        _issue_template_section(doc, issue.issue_no, issue.title, accts, names, listed=True)

    doc.add_page_break()
    _heading(doc, "【비상장】 한국공인회계사회 4대 중점", 1)
    for issue in fss_focus._DEFAULT_UNLISTED_2026:
        accts = ", ".join(issue.related_accounts) if issue.related_accounts else ""
        names = [c.name for c in issue.checklist] or [""]
        _issue_template_section(doc, issue.issue_no, issue.title, accts, names, listed=False)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_focus_checklist_templates_xlsx(*, is_listed: bool) -> bytes:
    """4대 중점 xlsx (지적사례 체크리스트 — Hanul DB 배포용)."""
    import focus_checklist_export as fce

    return fce.build_focus_checklist_workbook(is_listed=is_listed)


def _styled_sheet(wb, title: str, headers: list[str], widths: list[int]):
    """헤더 스타일 적용된 워크시트 생성."""
    from openpyxl.styles import Font, PatternFill
    from openpyxl.utils import get_column_letter

    ws = wb.active
    ws.title = title
    hdr_fill = PatternFill("solid", fgColor="1F3864")
    for j, h in enumerate(headers, 1):
        c = ws.cell(1, j, h)
        c.fill = hdr_fill
        c.font = Font(name=_FONT, bold=True, color="FFFFFF", size=10)
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    return ws


def build_sheet_tieout_template_xlsx() -> bytes:
    """§D 조서 연결·대사 사전 xlsx."""
    from openpyxl import Workbook

    headers = [
        "sheet_code",
        "role",
        "compare_unit",
        "exclude_from_cross_tieout",
        "related_accounts",
        "notes",
    ]
    wb = Workbook()
    ws = _styled_sheet(wb, "조서대사", headers, [12, 22, 18, 10, 20, 28])
    samples = [
        ("BB100", "차입금_잔액_유형별", "유형합산 또는 합계행", "N", "차입금", "단기+유동성+장기 합계"),
        ("BB200", "차입금_이자테스트", "해당없음", "Y", "차입금", "평균이자·이자비용 테스트 — 교차합 제외"),
        ("BB Lead", "차입금_리드", "Lead 행 당기금액", "N", "차입금", ""),
        ("CL100", "우발부채_제공보증", "조회서↔주석", "N", "우발부채·약정", "제공/수혜 보증 분리"),
        ("A Lead", "현금_리드", "Lead 행 당기금액", "N", "현금및예금", ""),
        ("F200", "투자자산_세부", "시트별 판단", "N", "투자자산", "손상·지분법 검토"),
    ]
    for i, row in enumerate(samples, 2):
        for j, v in enumerate(row, 1):
            ws.cell(i, j, v)
    for i in range(len(samples) + 2, len(samples) + 12):
        ws.cell(i, 1, "")
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_review_phrases_template_xlsx() -> bytes:
    """§E 검토내역·결론 인정 문장 사전 xlsx."""
    from openpyxl import Workbook

    headers = [
        "account",
        "topic",
        "acceptable_phrases",
        "adjacent_cell_ok",
        "red_flag_phrases",
        "context_hint",
        "notes",
    ]
    wb = Workbook()
    ws = _styled_sheet(wb, "인정문장", headers, [14, 22, 36, 8, 28, 16, 24])
    samples = [
        (
            "투자자산",
            "손상·회수가능성",
            "회수가능;추가 대손충당금은 반영하지 않음;손상 없음;공정가치 적정",
            "Y",
            "손상 징후 미검토;추가 충당 필요",
            "결론셀·각주",
            "F200 C42-C43 유형",
        ),
        (
            "투자자산",
            "지분법·손상·공시",
            "지분법;손상검토;회수가능;손상기준;평가·공시",
            "Y",
            "",
            "[검토 요청] 투자자산 bullet",
            "5개 bullet 통합 해소",
        ),
        (
            "매출채권",
            "ECL·대손",
            "기대신용손실;12개월;전체기간;회수가능;손실충당금",
            "Y",
            "충당 미반영",
            "",
            "",
        ),
        (
            "차입금",
            "조회·대사",
            "조회 결과 일치;True;대사 완료;조회서",
            "Y",
            "조회 미수행",
            "Lead·BB100",
            "",
        ),
        (
            "재고자산",
            "저가법·NRV",
            "순실현가능가치;저가법;평가손실;NRV",
            "Y",
            "평가손실 미인식",
            "",
            "",
        ),
    ]
    for i, row in enumerate(samples, 2):
        for j, v in enumerate(row, 1):
            ws.cell(i, j, v)
    for i in range(len(samples) + 2, len(samples) + 15):
        ws.cell(i, 1, "")
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_golden_set_template_xlsx() -> bytes:
    """§H 골든셋·회귀 기준 xlsx (4대 중점 체크리스트 연동)."""
    import golden_set_export as gse

    return gse.build_golden_set_workbook()


def build_procedure_catalog_template_xlsx() -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "필수절차"
    headers = [
        "account",
        "procedure_id",
        "procedure_name",
        "trigger",
        "required_evidence",
        "detect_all",
        "detect_any",
        "acceptable_conclusion",
        "cross_sheet",
        "fail_if",
        "basis",
        "to_be",
    ]
    hdr_fill = PatternFill("solid", fgColor="1F3864")
    for j, h in enumerate(headers, 1):
        c = ws.cell(1, j, h)
        c.fill = hdr_fill
        c.font = Font(name=_FONT, bold=True, color="FFFFFF", size=10)
    sample = [
        ("차입금", "BB-01", "금융기관 외부조회", "잔액>PM", "조회서;대사;조회일", "조회", "confirmation", "일치", "BB100;Lead", "Ref만", "KSA505", ""),
        ("현금및예금", "A-01", "예금 조회", "잔액>PM", "조회서;tick", "조회", "", "", "A Lead", "", "KSA505", ""),
    ]
    for i, row in enumerate(sample, 2):
        for j, v in enumerate(row, 1):
            ws.cell(i, j, v)
    for i in range(4, 14):
        ws.cell(i, 1, "")
    for i, w in enumerate([14, 10, 22, 14, 24, 16, 16, 20, 16, 18, 14, 24], 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def ensure_local_templates() -> Path:
    """data/templates/ 에 빈 양식 xlsx 생성."""
    root = Path(__file__).resolve().parent / "data" / "templates"
    root.mkdir(parents=True, exist_ok=True)
    files = {
        rg.TEMPLATE_FILES["focus_listed"]: lambda: build_focus_checklist_templates_xlsx(is_listed=True),
        rg.TEMPLATE_FILES["focus_unlisted"]: lambda: build_focus_checklist_templates_xlsx(is_listed=False),
        rg.TEMPLATE_FILES["procedures"]: build_procedure_catalog_template_xlsx,
        rg.TEMPLATE_FILES["sheet_tieout"]: build_sheet_tieout_template_xlsx,
        rg.TEMPLATE_FILES["review_phrases"]: build_review_phrases_template_xlsx,
        rg.TEMPLATE_FILES["golden_set"]: build_golden_set_template_xlsx,
    }
    for name, builder in files.items():
        p = root / name
        if not p.is_file():
            p.write_bytes(builder())
    return root
