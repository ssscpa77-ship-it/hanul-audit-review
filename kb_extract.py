"""지식 저장소 문서 텍스트 추출기.

Hanul DB 안의 다양한 형식(PDF·DOCX·XLSX·HWPX·HWP)에서 본문 텍스트를
추출해, 색인(build_index.py)에 사용할 수 있게 정규화합니다.

스캔 PDF: 내장 텍스트가 없으면 Tesseract OCR(kor+eng) 폴백.
  brew install tesseract tesseract-lang
  HANUL_OCR=0 으로 OCR 비활성화 가능.
"""

from __future__ import annotations

import hashlib
import os
import re
import shutil
import unicodedata
import zipfile

SUPPORTED_EXTS = {"pdf", "docx", "xlsx", "xls", "xlsm", "hwpx", "hwp"}

_OCR_MIN_CHARS = 30
_OCR_MAX_PAGES = 80
_OCR_SCALE = 2.0
_OCR_LANG = os.environ.get("HANUL_OCR_LANG", "kor+eng")
_OCR_CACHE_DIR = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "kb_store", "ocr_cache"
)

# 대용량 감사조서 xlsx/xls — 시트·행 상한 (용량 제한이 아니라 추출 안정성)
_EXCEL_MAX_SHEETS = 50
_EXCEL_MAX_ROWS = 1000
_EXCEL_MAX_COLS = 50
_EXCEL_MAX_CHARS = 2_000_000


def ext_of(path: str) -> str:
    return path.rsplit(".", 1)[-1].lower() if "." in path else ""


def extract_text(path: str) -> str:
    """파일 형식에 맞는 추출기로 본문 텍스트를 반환. 실패 시 빈 문자열."""
    ext = ext_of(path)
    try:
        if ext == "pdf":
            return _pdf(path)
        if ext == "docx":
            return _docx(path)
        if ext in ("xlsx", "xls", "xlsm"):
            return _excel(path)
        if ext == "hwpx":
            return _hwpx(path)
        if ext == "hwp":
            return _hwp(path)
    except Exception:  # noqa: BLE001 - 개별 파일 실패는 건너뜀
        return ""
    return ""


def ocr_available() -> bool:
    """Tesseract OCR 사용 가능 여부."""
    if os.environ.get("HANUL_OCR", "1").strip().lower() in ("0", "false", "no"):
        return False
    return shutil.which("tesseract") is not None


def _pdf(path: str) -> str:
    text = _pdf_text_layer(path)
    if len(normalize(text)) >= _OCR_MIN_CHARS:
        return text
    if ocr_available():
        ocr_text = _pdf_ocr(path)
        if len(normalize(ocr_text)) >= _OCR_MIN_CHARS:
            return ocr_text
    return text


def _pdf_text_layer(path: str) -> str:
    import pypdfium2 as pdfium

    parts: list[str] = []
    pdf = pdfium.PdfDocument(path)
    try:
        for page in pdf:
            textpage = page.get_textpage()
            parts.append(textpage.get_text_range() or "")
            textpage.close()
            page.close()
    finally:
        pdf.close()
    return "\n".join(parts)


def _ocr_cache_key(path: str) -> str:
    try:
        st = os.stat(path)
        raw = f"{os.path.abspath(path)}:{st.st_mtime_ns}:{st.st_size}"
    except OSError:
        raw = os.path.abspath(path)
    return hashlib.sha256(raw.encode()).hexdigest()[:32]


def _ocr_cache_read(path: str) -> str:
    key = _ocr_cache_key(path)
    cache_path = os.path.join(_OCR_CACHE_DIR, f"{key}.txt")
    if not os.path.isfile(cache_path):
        return ""
    try:
        with open(cache_path, encoding="utf-8") as f:
            return f.read()
    except OSError:
        return ""


def _ocr_cache_write(path: str, text: str) -> None:
    os.makedirs(_OCR_CACHE_DIR, exist_ok=True)
    key = _ocr_cache_key(path)
    cache_path = os.path.join(_OCR_CACHE_DIR, f"{key}.txt")
    try:
        with open(cache_path, "w", encoding="utf-8") as f:
            f.write(text)
    except OSError:
        pass


def _pdf_ocr(path: str) -> str:
    """스캔 PDF — 페이지 렌더 후 Tesseract OCR."""
    cached = _ocr_cache_read(path)
    if cached:
        return cached

    import pytesseract

    pytesseract.pytesseract.tesseract_cmd = shutil.which("tesseract") or "tesseract"

    import pypdfium2 as pdfium

    parts: list[str] = []
    pdf = pdfium.PdfDocument(path)
    try:
        n_pages = len(pdf)
        limit = min(n_pages, _OCR_MAX_PAGES)
        for i in range(limit):
            page = pdf[i]
            try:
                bitmap = page.render(scale=_OCR_SCALE)
                pil = bitmap.to_pil()
                try:
                    txt = pytesseract.image_to_string(pil, lang=_OCR_LANG)
                    if txt.strip():
                        parts.append(txt)
                finally:
                    pil.close()
                    bitmap.close()
            finally:
                page.close()
        if n_pages > _OCR_MAX_PAGES:
            parts.append(f"\n[OCR: {_OCR_MAX_PAGES}/{n_pages}페이지까지만 처리]")
    finally:
        pdf.close()

    text = "\n".join(parts)
    if text.strip():
        _ocr_cache_write(path, text)
    return text


def _docx(path: str) -> str:
    from docx import Document

    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _excel(path: str) -> str:
    """xlsx/xlsm/xls — openpyxl·xlrd 기반 (대용량·혼합형 셀 안전)."""
    ext = ext_of(path)
    if ext == "xls":
        return _excel_xlrd(path)
    if ext in ("xlsx", "xlsm"):
        return _excel_openpyxl(path)
    return ""


def _cell_str(val) -> str:
    if val is None:
        return ""
    if isinstance(val, bool):
        return "TRUE" if val else "FALSE"
    if isinstance(val, float):
        if val == int(val) and abs(val) < 1e15:
            return str(int(val))
    s = str(val).strip()
    if s.lower() in ("nan", "none"):
        return ""
    return s


def _excel_openpyxl(path: str) -> str:
    from openpyxl import load_workbook

    parts: list[str] = []
    total = 0
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        sheets = [
            ws for ws in wb.worksheets
            if getattr(ws, "sheet_state", "visible") == "visible"
        ] or list(wb.worksheets)
        for ws in sheets[:_EXCEL_MAX_SHEETS]:
            header = f"[시트: {ws.title}]"
            parts.append(header)
            total += len(header)
            lines: list[str] = []
            for ri, row in enumerate(ws.iter_rows(values_only=True)):
                if ri >= _EXCEL_MAX_ROWS:
                    lines.append(f"[행 {_EXCEL_MAX_ROWS}행까지 추출]")
                    break
                cells = [_cell_str(c) for c in row[:_EXCEL_MAX_COLS]]
                line = " | ".join(c for c in cells if c)
                if line:
                    lines.append(line)
            block = "\n".join(lines)
            if block:
                parts.append(block)
                total += len(block)
            if total >= _EXCEL_MAX_CHARS:
                parts.append("[문서 텍스트 상한 도달]")
                break
    finally:
        wb.close()
    return "\n".join(parts)[:_EXCEL_MAX_CHARS]


def _excel_xlrd(path: str) -> str:
    import xlrd

    parts: list[str] = []
    total = 0
    bk = xlrd.open_workbook(path, on_demand=True)
    try:
        sheet_indexes = [
            i for i in range(bk.nsheets)
            if bk.sheet_by_index(i).visibility == 0
        ] or list(range(bk.nsheets))
        for si in sheet_indexes[:_EXCEL_MAX_SHEETS]:
            sh = bk.sheet_by_index(si)
            header = f"[시트: {sh.name}]"
            parts.append(header)
            total += len(header)
            lines: list[str] = []
            for r in range(min(sh.nrows, _EXCEL_MAX_ROWS)):
                cells = []
                for c in range(min(sh.ncols, _EXCEL_MAX_COLS)):
                    v = _cell_str(sh.cell_value(r, c))
                    if v:
                        cells.append(v)
                if cells:
                    lines.append(" ".join(cells))
            block = "\n".join(lines)
            if block:
                parts.append(block)
                total += len(block)
            if total >= _EXCEL_MAX_CHARS:
                parts.append("[문서 텍스트 상한 도달]")
                break
    finally:
        bk.release_resources()
    return "\n".join(parts)[:_EXCEL_MAX_CHARS]


_HP_T_RE = re.compile(r"<hp:t[^>]*>(.*?)</hp:t>", re.DOTALL)
_TAG_RE = re.compile(r"<[^>]+>")


def _hwpx(path: str) -> str:
    parts: list[str] = []
    with zipfile.ZipFile(path) as zf:
        names = sorted(
            n for n in zf.namelist() if re.search(r"Contents/section\d+\.xml$", n)
        )
        for name in names:
            xml = zf.read(name).decode("utf-8", errors="ignore")
            for m in _HP_T_RE.findall(xml):
                text = _TAG_RE.sub("", m)
                if text.strip():
                    parts.append(text)
    return "\n".join(parts)


def _hwp(path: str) -> str:
    """HWP 5.0(OLE) 본문 텍스트 추출. 실패 시 미리보기 텍스트로 대체."""
    import zlib

    import olefile

    ole = olefile.OleFileIO(path)
    try:
        # 압축 여부 확인 (FileHeader 37번째 바이트 bit0)
        compressed = True
        if ole.exists("FileHeader"):
            header = ole.openstream("FileHeader").read()
            if len(header) > 36:
                compressed = bool(header[36] & 0x01)

        sections = sorted(
            (e for e in ole.listdir() if len(e) == 2 and e[0] == "BodyText"),
            key=lambda e: e[1],
        )
        out: list[str] = []
        for entry in sections:
            data = ole.openstream(entry).read()
            if compressed:
                try:
                    data = zlib.decompress(data, -15)
                except Exception:  # noqa: BLE001
                    continue
            out.append(_hwp_records_to_text(data))
        text = "\n".join(p for p in out if p.strip())
        if text.strip():
            return text

        # 본문 실패 시 미리보기 텍스트(PrvText, UTF-16LE)
        if ole.exists("PrvText"):
            return ole.openstream("PrvText").read().decode("utf-16-le", errors="ignore")
        return ""
    finally:
        ole.close()


# HWPTAG_PARA_TEXT = HWPTAG_BEGIN(0x10) + 51
_HWPTAG_PARA_TEXT = 0x10 + 51


def _hwp_records_to_text(data: bytes) -> str:
    """HWP BodyText 레코드 스트림에서 문단 텍스트(tag 67)만 추출."""
    parts: list[str] = []
    pos, size = 0, len(data)
    while pos + 4 <= size:
        header = int.from_bytes(data[pos : pos + 4], "little")
        tag_id = header & 0x3FF
        rec_size = (header >> 20) & 0xFFF
        pos += 4
        if rec_size == 0xFFF:
            if pos + 4 > size:
                break
            rec_size = int.from_bytes(data[pos : pos + 4], "little")
            pos += 4
        payload = data[pos : pos + rec_size]
        pos += rec_size
        if tag_id == _HWPTAG_PARA_TEXT:
            parts.append(_decode_hwp_text(payload))
    return "".join(parts)


def _decode_hwp_text(payload: bytes) -> str:
    """UTF-16LE 문단 페이로드에서 제어문자를 건너뛰고 텍스트만 추출."""
    out: list[str] = []
    i, n = 0, len(payload)
    while i + 2 <= n:
        code = int.from_bytes(payload[i : i + 2], "little")
        i += 2
        if code in (10, 13):
            out.append("\n")
        elif code < 32:
            # 확장 제어문자(일부는 인라인 14바이트 추가) 처리
            if code in (1, 2, 3, 4, 5, 11, 12, 14, 15, 16, 17, 18, 21, 22, 23):
                i += 14
            # 그 외 제어문자는 무시
        else:
            out.append(chr(code))
    return "".join(out)


_WS_RE = re.compile(r"[ \t\u00a0]+")
_NL_RE = re.compile(r"\n{3,}")
# 잘못된 대리 문자(surrogate) 및 널 문자 — SQLite/UTF-8 저장 불가
_BAD_RE = re.compile(r"[\ud800-\udfff\x00]")


def normalize(text: str) -> str:
    text = unicodedata.normalize("NFC", text)
    text = _BAD_RE.sub("", text)
    text = _WS_RE.sub(" ", text)
    text = _NL_RE.sub("\n\n", text)
    return text.strip()


def chunk_text(text: str, size: int = 800, overlap: int = 120) -> list[str]:
    """정규화된 텍스트를 겹침이 있는 청크로 분할."""
    text = normalize(text)
    if not text:
        return []
    chunks: list[str] = []
    start, n = 0, len(text)
    while start < n:
        end = min(start + size, n)
        # 문장/줄 경계에서 자르기 시도
        if end < n:
            boundary = text.rfind("\n", start + size // 2, end)
            if boundary == -1:
                boundary = text.rfind(". ", start + size // 2, end)
            if boundary != -1 and boundary > start:
                end = boundary + 1
        chunk = text[start:end].strip()
        if len(chunk) >= 30:
            chunks.append(chunk)
        if end >= n:
            break
        start = max(end - overlap, start + 1)
    return chunks
