"""한울 DB 「한울 개발지침서」 협업용 문서 생성."""

from __future__ import annotations

from datetime import date
from io import BytesIO
from pathlib import Path

import knowledge_base as kb
import review_guidelines as rg

DEV_GUIDE_DIR = "한울 개발지침서"
_FONT = "맑은 고딕"


def hanul_dev_guide_root() -> Path:
    return Path(kb.SOURCE_DIR) / DEV_GUIDE_DIR


def guide_basename(d: date | None = None) -> str:
    d = d or date.today()
    return f"개발지침서_한울_{d.strftime('%Y-%m-%d')}"


def _kb_stats_line() -> str:
    try:
        s = kb.stats()
        return f"{s['documents']:,}문서 / {s['chunks']:,}청크"
    except Exception:  # noqa: BLE001
        return "(색인 미생성 — build_index.py 실행 필요)"


def build_development_guide_markdown(d: date | None = None) -> str:
    """협업용 개발지침서 본문 (Markdown)."""
    d = d or date.today()
    ds = d.strftime("%Y-%m-%d")
    kb_stats = _kb_stats_line()
    return f"""# 한울 감사조서 자가검토 — 개발지침서

> **문서명:** 개발지침서_한울_{ds}  
> **프로젝트:** hanul-002 (Streamlit 자가 심리 에이전트)  
> **목적:** 직원 공유·개발 진행 체크·협업 자료 허브 (living document)  
> **최종 갱신:** {ds}  
> **앱 버전:** v0.8.0 · FY2026 MVP

---

## 목차

1. [한 줄 정의](#1-한-줄-정의)  
2. [Hanul DB 폴더 구조](#2-hanul-db-폴더-구조)  
3. [검토 파이프라인](#3-검토-파이프라인)  
4. [QC 3축·Tier 우선순위](#4-qc-3축tier-우선순위)  
5. [핵심 수정지시 요약 (§0)](#5-핵심-수정지시-요약-0)  
6. [Hanul DB KB(RAG) 학습](#6-hanul-db-kbrag-학습)  
7. [감리지적 체크리스트](#7-감리지적-체크리스트)  
8. [표준절차·리뷰노트 KB 연동](#8-표준절차리뷰노트-kb-연동)  
9. [4대 중점·자가검토 템플릿](#9-4대-중점자가검토-템플릿)  
10. [골든셋 회귀](#10-골든셋-회귀)  
11. [배포·접속 URL](#11-배포접속-url)  
12. [환경 변수](#12-환경-변수)  
13. [실행·검증 명령어](#13-실행검증-명령어)  
14. [셸 스크립트 목록](#14-셸-스크립트-목록)  
15. [핵심 모듈 맵](#15-핵심-모듈-맵)  
16. [3층 엔진·로드맵](#16-3층-엔진로드맵)  
17. [알려진 한계·미완료](#17-알려진-한계미완료)  
18. [다음 스프린트](#18-다음-스프린트)  
19. [참고 문서](#19-참고-문서)

---

## 1. 한 줄 정의

회계법인 QRM 제출 전, 감사조서를 **규칙엔진 + Hanul DB(RAG) + AI(선택)** 로 자가검토하여  
**출처가 명확한 리뷰노트**를 생성하는 시스템.

| 항목 | 내용 |
|------|------|
| 로컬 URL | http://localhost:8505 |
| 사내망 URL | `config.primary_access_url()` (서버 IP:8505) |
| 외부 검증 URL | Cloudflare Quick Tunnel + 공유 랜딩 (`share_gateway.py` :8506) |
| 소스 경로 | `ABC05CEO/hanul-002/` |
| Hanul DB | `{kb.SOURCE_DIR}` |
| KB 색인 | `kb_store/hanul_kb.sqlite` — **{kb_stats}** |

---

## 2. Hanul DB 폴더 구조

```
Hanul DB/
├── 4대 중점사항 감리대상/          … 금감원·한공회 원문 PDF
├── 금융감독원 감리지적사례/        … FSS 감리지적 (KB·체크리스트)
├── 한국공인회계사회 감리지적사례/  … KICPA 감리지적
├── 회계감사기준/                   … KSA
├── 한국채택국제회계기준/           … K-IFRS
├── 일반기업회계기준/               … K-GAAP
├── K-IFRS 실무사례와 해설/
├── 질의회신_한국회계기준원/        … 기준원 질의회신
├── 한울 예시조서/                  … 4000 실증절차·Audit Program
├── 한울 예시보고서/
├── 감사조서 샘플/                  … 실무 조서·엑셀 샘플 (163건+)
├── 자가검토_지침_템플릿/           … xlsx 체크리스트·인정문장
│   ├── 4대중점_체크리스트_상장/비상장_FY2026.xlsx
│   ├── 감리지적_체크리스트_상장/비상장_FY2026.xlsx
│   └── (기타 템플릿 6종)
└── 한울 개발지침서/               … ★ 본 문서·진행상황·회의록
    ├── 개발지침서_한울_{ds}.docx / .md
    ├── README.md
    ├── 진행상황/
    │   ├── 개발진행_체크리스트_FY2026.xlsx
    │   └── 변경이력.md
    └── 회의록/
```

**협업 규칙**
- QC·운영: `자가검토_지침_템플릿/` xlsx 작성·갱신
- 개발: 코드 연동 후 `build_index.py` · `golden_set_regression.py` 확인
- 중요 결정: `진행상황/변경이력.md` 또는 `회의록/` 기록

---

## 3. 검토 파이프라인

### 3.1 전체 흐름 (`app.run_analysis`)

```
업로드(parse_uploads)
  → run_review (규칙엔진)
  → [AI ON] ai_review.run_sheet_reviews
  → note_merge.merge_review_notes
  → notes_pipeline.post_process_notes
  → _attach_references (Hanul DB 근거·감리사례·조서예시)
  → 대시보드·엑셀 출력
```

### 3.2 `run_review` 내부 순서

| 순서 | 모듈 | 역할 |
|------|------|------|
| 1 | `_check_procedures` | 축1: 핵심 감사절차 누락 (KB + PROCEDURE_RULES) |
| 2 | `fss_focus.run_focus_review` | 축2: 4대 중점 체크리스트 |
| 3 | `enforcement_review` | 감리지적 체크리스트 (조서인덱스별 필수) |
| 4 | 계산·주석·대사·우발·교차 FS | `review_engine` |
| 5 | `qc_review.run_qc_checks` | 품질관리 |
| 6 | 중요도 「하」 필터 | focus/enforcement 보호 노트 제외 |

### 3.3 `post_process_notes` 후처리

- `filter_off_account_notes` — 조서에 없는 계정 지적 제거
- 교차조서: 절차·공시·현금 외부조회 중복 제거
- `note_merge` — TUL·우발·차입금 담보 등 통합
- `qc_review.prune_documented_review_notes`
- `prune_qc_notes` — 경미·형식 노트 제거
- `simplify_note_outputs` — basis 한 줄화

---

## 4. QC 3축·Tier 우선순위

| Tier | 성격 | 카테고리 | 중요도 |
|------|------|----------|--------|
| **1** | 4대 중점사항 | `중점감리` | 상 (`focus_protected`) |
| **2** | 핵심 절차 누락 | `절차누락` | 상 |
| **3** | 유사 감리지적 리스크 | 기존 노트 ⚖️ 보강 | — |
| — | 경미·형식 | 형식·완전성 등 | 하 (기본 제외) |

**3축**
1. **절차** — Hanul DB 표준조서·실무조서 기준 `_check_procedures`
2. **4대 중점** — 금감원(상장)·한공회(비상장) 체크리스트 `fss_focus`
3. **감리사례** — `_attach_references` · `gather_citations` (맥락·계정 일치 시만)

---

## 5. 핵심 수정지시 요약 (§0)

> 상세: `docs/자가검토_작업지시서_전체설정.md`

| 번호 | 원칙 | 요약 |
|------|------|------|
| 0.1 | 조서 계정만 리뷰 | `sheet_code_registry` 조서번호 최우선; 없는 계정 지적 금지 |
| 0.2 | 중요한 누락만 | 중요도 「하」·형식 검사 기본 제외; AI 시트당 최대 2건 |
| 0.3 | 숨김 시트 제외 | visible 시트만; 모든 노트에 시트명 표기 |
| 0.4 | Lead·주석 통합시트 | tick mark≠레퍼런스; 단위·절사 허용; BB 교차합 총합 기준 |
| 0.5 | 4대 중점 | 상장/비상장·K-IFRS/K-GAAP 선택; checklist_id당 1건 |
| 0.6 | 리뷰근거 간결 | 기준서 한 줄; 감리사례 ⚖️만; 맥락 2키워드 이상 일치 |
| 0.7 | 회사명 | Lead 시트 상단 최빈값; 타 법인명 혼입 금지 |
| 0.8 | A Lead 현금조회 | tick·Ref·조회서 → 외부조회 수행 인정 |
| 0.9 | 감리사례 | 노트 있을 때만; 계정·주제 일치만 |
| 0.10 | TUL 부외부채 | 유사 지적 1건 통합·리마인드형 |
| 0.11 | QC 3축 | §4 참조 |

---

## 6. Hanul DB KB(RAG) 학습

### 6.1 색인 빌드

```bash
cd hanul-002
VENV/bin/python build_index.py              # 전체 색인
VENV/bin/python build_index.py --retry-empty  # empty 문서만 재색인(OCR·엑셀)
VENV/bin/python build_index.py --reset      # DB 초기화 후 전체
```

### 6.2 KB 카테고리 (`knowledge_base.py`)

| 상수 | Hanul DB 폴더 |
|------|----------------|
| `STANDARDS_CATEGORIES` | 회계감사기준, K-IFRS, K-GAAP, K-IFRS 실무해설 |
| `QNA_CATEGORIES` | 질의회신_한국회계기준원 |
| `ENFORCEMENT_CATEGORIES` | 금감원·한공회 감리지적사례 |
| `WORKPAPER_CATEGORIES` | 한울 예시조서, 한울 예시보고서, **감사조서 샘플** |
| `FOCUS_CATEGORIES` | 4대 중점사항 감리대상 |
| `GUIDELINES_CATEGORIES` | 자가검토_지침_템플릿 |
| `REVIEW_KB_CATEGORIES` | 위 전체 통합 (리뷰 근거 검색) |

### 6.3 OCR (스캔 PDF)

- **모듈:** `kb_extract.py` — Tesseract `kor+eng` 폴백
- **캐시:** `kb_store/ocr_cache/`
- **비활성:** `HANUL_OCR=0`
- **의존:** `brew install tesseract tesseract-lang`; `pytesseract`, `Pillow`
- **재색인:** `build_index.py --retry-empty`

### 6.4 엑셀 추출 (감사조서 샘플)

- **원인(2026-07):** pandas `.str.cat` 혼합형 셀 오류 (용량 아님)
- **해결:** openpyxl `read_only` (xlsx/xlsm), xlrd (xls)
- **상한:** 시트 50개, 행 1000, 열 50, 총 2M자
- **결과:** 감사조서 샘플 163건·한울 예시조서 150건 복구

### 6.5 `gather_citations` (AI·근거)

| 그룹 | k (기본) | 용도 |
|------|----------|------|
| 기준 | 3 | KSA·K-IFRS |
| 질의회신 | 2 | 기준원 QnA |
| 감리지적사례 | 2 | FSS·KICPA |
| 감사절차 예시 | **3** | WORKPAPER |
| 4대중점 | 1 | FOCUS |

### 6.6 학습 현황 ({ds} 기준)

| 구분 | 상태 |
|------|------|
| K-IFRS, K-GAAP, KSA, 실무해설, 감리지적 479건 | ✅ |
| 4대중점, 자가검토 템플릿, 예시보고서 | ✅ |
| 질의회신 2,074건 | ✅ |
| 감사조서 샘플·한울 예시조서·실제 조서 엑셀 | ✅ |
| **질의회신 빈 PDF 74건** | ❌ 원본 PDF 교체 필요 |

---

## 7. 감리지적 체크리스트

### 7.1 개요

- **모듈:** `enforcement_review.py`, `enforcement_case_learner.py`
- **로더:** `guidelines_loader.load_enforcement_checklist_from_db()`
- **배포:** `python enforcement_checklist_export.py` → Hanul DB xlsx
- **상세화:** `case_context`, `audit_focus`, `case_narratives` (학습·분류)

### 7.2 실행

- `run_enforcement_checklist_review()` — `run_review` 내 4대 중점 직후
- 조서인덱스(`sheet_code`)별 **전 항목 필수 검토**
- gap: `검토누락` / `결론미비` → category `감리지적체크`, importance `중`
- `enforcement_protected=True` — QC prune 보호

### 7.3 xlsx 위치

`자가검토_지침_템플릿/감리지적_체크리스트_상장_FY2026.xlsx`  
`자가검토_지침_템플릿/감리지적_체크리스트_비상장_FY2026.xlsx`

---

## 8. 표준절차·리뷰노트 KB 연동

### 8.1 `get_standard_procedures` (2026-07 보강)

- **검색:** 계정별 절차 쿼리 + 조서번호(E,C 등) + 동의어
- **매칭:** `_matches_workpaper_account` — 스니펫·조서코드·동의어
- **라벨:** 파일명 대신 의미 라벨 (`재고자산 실사입회`, `대손충당금 검토` 등)
- **필터:** 스프레드시트 덤프·숫자 위주 스니펫 제외; 계정 간 라벨 오염 방지
- **연동:** `review_engine._check_procedures` — 동의어 전달, `kb_snippet` reason 인용

### 8.2 `_attach_references` 보강

- `WORKPAPER_CATEGORIES` basis·`workpaper_ref` (절차누락)
- 기준·질의회신·4대중점·**조서예시** 통합 검색
- 계정·주제 맥락 필터 (`is_off_account_text`, `_is_relevant`)

### 8.3 FILELIKE 필터

- KB 파일명(`_FY2026_v1` 등)만 있는 항목은 절차명에서 제외
- 의미 있는 절차 라벨은 유지

---

## 9. 4대 중점·자가검토 템플릿

### 9.1 연동 현황 (`{rg.GUIDELINES_DB_SUBDIR}`)

| 템플릿 | 연동 | 비고 |
|--------|------|------|
| 4대중점_체크리스트_상장/비상장 | ✅ `fss_focus` | 17/22항목 |
| 감리지적_체크리스트_상장/비상장 | ✅ `enforcement_review` | case_context 상세화 |
| 조서연결_대사_사전 | ✅ `review_engine` | BB·CL·F 등 |
| 계정별_필수절차_카탈로그 | 🔶 부분 | KB `get_standard_procedures` |
| 검토내역_결론_인정문장 | ⏳ | qc_review 연동 예정 |
| 골든셋_회귀기준 | ✅ | `golden_set_regression.py` 11/11 |

### 9.2 배포 명령

```bash
python focus_checklist_export.py
python enforcement_checklist_export.py
```

---

## 10. 골든셋 회귀

```bash
VENV/bin/python golden_set_regression.py
```

| 케이스 | 내용 |
|--------|------|
| GS-F-L2-01~06 | 재고 저가법·NRV·진부화 (상장) |
| GS-F-U2-01~06 | 지분법 (비상장) |
| GS-F-L2-PASS / GS-F-U2-PASS | 충분 조서 → 0건 |
| GS-001 | 양호 조서 Tier1 0건 |

**현재:** **11/11 통과** ({ds} 기준)

---

## 11. 배포·접속 URL

### 11.1 로컬·사내망

```bash
./run_server.sh
# → http://localhost:8505
# → http://{{LAN_IP}}:8505  (동일 Wi-Fi)
```

- `.streamlit/config.toml`: `address = "0.0.0.0"`, `port = 8505`
- `config.deployment_access_urls()` / `primary_access_url()`

### 11.2 외부 검증 (교수님·원격)

**2터널 구조**

| 포트 | 서비스 | URL 파일 |
|------|--------|----------|
| 8505 | Streamlit 앱 | `.app_tunnel.url` |
| 8506 | 공유 랜딩 (`share_gateway.py`) | `.tunnel.url` ← **카톡 공유용** |

```bash
./run_tunnels.sh          # 앱+공유 터널 동시
./keep_alive.sh           # 서버·게이트웨이·터널 상시 유지
./disable_sleep.sh        # 절전 완화 (검증 기간)
```

**공유 랜딩 (`share/share_landing.html`)**
- 제목: **ABC 5기 신성섭 · 감사조서 Smart Reviewer**
- 디자인: 진한 파란 배경 + 흰 글씨 + OG 이미지 (`share/share_og.png`)
- 카톡 메시지 템플릿: `share/kakao_message.txt` → `.kakao_message_ready.txt`

### 11.3 도메인 연결 (추후)

`.env` → `APP_BASE_URL=https://your-domain` 설정 시 IP 대체

---

## 12. 환경 변수 (`.env`)

| 변수 | 설명 |
|------|------|
| `OPENAI_API_KEY` | AI 심층리뷰 |
| `HANUL_DB_PATH` | Hanul DB 경로 |
| `FSS_FOCUS_ISSUES_YEAR` | 4대 중점 연도 (2026) |
| `FSS_FOCUS_ISSUES_LISTED` | 상장 4대 이슈 목록 |
| `FALLBACK_TO_HARDCODED_PROCEDURES` | KB 없을 때 PROCEDURE_RULES |
| `APP_BASE_URL` | 사내/도메인 URL (비우면 IP 자동) |
| `STREAMLIT_SERVER_PORT` | 기본 8505 |
| `HANUL_OCR` | 0=OCR 비활성 |

---

## 13. 실행·검증 명령어

```bash
cd hanul-002
source VENV/bin/activate

# 앱
./run_server.sh

# KB
VENV/bin/python build_index.py
VENV/bin/python build_index.py --retry-empty

# 체크리스트 배포
VENV/bin/python focus_checklist_export.py
VENV/bin/python enforcement_checklist_export.py

# 회귀
VENV/bin/python golden_set_regression.py

# 개발지침서 배포 (Hanul DB)
VENV/bin/python dev_guide_export.py

# KB 상태
VENV/bin/python -c "import knowledge_base as kb; print(kb.stats())"
```

---

## 14. 셸 스크립트 목록

| 스크립트 | 용도 |
|----------|------|
| `run_server.sh` | Streamlit 사내망 서버 |
| `run_tunnel.sh [port] [share|app]` | Cloudflare 터널 1개 |
| `run_tunnels.sh` | 앱(8505)+공유(8506) 터널 |
| `keep_alive.sh` | 상시 유지·자동 재시작 |
| `stop_keep_alive.sh` | keep_alive 중지 |
| `stop_tunnel.sh` | 터널 전체 중지 |
| `disable_sleep.sh` | 맥북 절전 완화 |
| `restore_sleep.sh` | 절전 설정 복원 |

---

## 15. 핵심 모듈 맵

| 파일 | 역할 |
|------|------|
| `app.py` | UI, `run_analysis`, `_attach_references` |
| `review_engine.py` | 규칙엔진·절차·교차합·주석 |
| `fss_focus.py` | 4대 중점 |
| `enforcement_review.py` | 감리지적 체크리스트 |
| `qc_review.py` | QC·검토내역 해소 |
| `ai_review.py` | AI 심층 (`gather_citations`) |
| `knowledge_base.py` | RAG·표준절차·인용 |
| `kb_extract.py` | PDF/OCR/엑셀 추출 |
| `build_index.py` | KB 색인 |
| `guidelines_loader.py` | xlsx 지침 로드 |
| `notes_pipeline.py` | 후처리·필터 |
| `sheet_code_registry.py` | FY2026 조서번호 맵 |
| `share_gateway.py` | 카톡 공유 랜딩 |
| `config.py` | 설정·배포 URL |
| `dev_guide_export.py` | 본 개발지침서 생성 |

---

## 16. 3층 엔진·로드맵

| 층 | 역할 | 상태 |
|----|------|------|
| L1 결정론 | 절차·대사·4대중점·감리지적 | ✅ 운영 |
| L2 증거 그래프 | Lead↔세부↔조회서 | 🔶 부분 |
| L3 AI 서술 | reason·to_be 보조 | ✅ 운영 |
| 학습 루프 | 골든셋 회귀 | ✅ 11/11 |

| Phase | 내용 | 상태 |
|-------|------|------|
| A | 4대중점·감리지적 체크리스트 | ✅ |
| B | KB 전체 학습·리뷰 연동 | ✅ ({ds}) |
| C | 외부 배포·공유 URL | ✅ (터널) |
| D | 검토내역 인정문장 xlsx 연동 | ⏳ |
| E | AI L1 Pass 시트 스킵 | ⏳ |

---

## 17. 알려진 한계·미완료

1. **질의회신 빈 PDF 74건** — 원본 재다운로드 후 `--retry-empty`
2. **터널 URL** — 재시작 시 변경됨; 검증 종료 후 `stop_tunnel.sh`
3. **키워드 절차 판정** — 표 구조·완전성 미검증
4. **감리사례 역방향 매칭** — 독립 Tier3 지적 미구현
5. **임베딩 검색** — FTS5 BM25만 (향후 확장)

---

## 18. 다음 스프린트

1. `검토내역_결론_인정문장` → `qc_review` 연동
2. 질의회신 74건 원본 보강·재색인
3. 고정 도메인 + HTTPS (`APP_BASE_URL`)
4. AI L1 Pass 시트 스킵 (비용 절감)
5. 주석 누락·전기대비 규칙 보강

---

## 19. 참고 문서

| 문서 | 위치 |
|------|------|
| 작업지시서 전체설정 | `docs/자가검토_작업지시서_전체설정.md` |
| 제품 기획서 | `docs/기획서.md` |
| 카톡 메시지 | `share/kakao_message.txt` |
| QC 자문의견 Word | 앱 다운로드 / `word_export.py` |

---

## 부록: 직원 역할 분담

| 역할 | 담당 |
|------|------|
| QC·감리 | 4대중점·감리지적·인정문장 xlsx |
| 조서·IT | 골든셋·조서연결 사전 |
| 개발 | 코드·KB·배포·회귀 |
| 파트너 | 골든셋 기대 지적 확정 |

---

*본 문서는 협업용 living document입니다. 변경 시 `진행상황/변경이력.md`에 요약을 남겨 주십시오.*  
*갱신 명령: `VENV/bin/python dev_guide_export.py`*
"""


def build_development_guide_docx(d: date | None = None) -> bytes:
    """개발지침서 Word."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    d = d or date.today()
    ds = d.strftime("%Y-%m-%d")
    md = build_development_guide_markdown(d)

    doc = Document()
    title = doc.add_heading("한울 감사조서 자가검토 — 개발지침서", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    r = p.add_run(f"문서명: 개발지침서_한울_{ds}  |  갱신: {ds}")
    r.font.size = Pt(10)

    for block in md.split("\n"):
        line = block.strip()
        if not line or line == "---":
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], 1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], 2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], 3)
        elif line.startswith("|") and "---" not in line:
            if line.startswith("| **") or line.startswith("| 파일") or line.startswith("| 항목"):
                continue
            cells = [c.strip() for c in line.split("|") if c.strip()]
            if len(cells) >= 2:
                doc.add_paragraph(
                    f"{cells[0]}: {cells[1]}"
                    + (f" — {cells[2]}" if len(cells) > 2 else "")
                )
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("```"):
            continue
        elif line.startswith(">"):
            doc.add_paragraph(line.lstrip("> ").strip())
        else:
            doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_progress_checklist_xlsx(d: date | None = None) -> bytes:
    """진행상황 체크리스트."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    from openpyxl.utils import get_column_letter

    d = d or date.today()
    ds = d.strftime("%Y-%m-%d")

    wb = Workbook()
    ws = wb.active
    ws.title = "개발진행"
    headers = [
        "ID",
        "영역",
        "항목",
        "상태",
        "담당",
        "목표일",
        "완료일",
        "비고",
    ]
    fill = PatternFill("solid", fgColor="1F3864")
    for j, h in enumerate(headers, 1):
        c = ws.cell(1, j, h)
        c.fill = fill
        c.font = Font(name=_FONT, bold=True, color="FFFFFF", size=10)

    rows = [
        ("D-01", "템플릿", "4대중점_상장 체크리스트", "완료", "", "", ds, "17항목 연동"),
        ("D-02", "템플릿", "4대중점_비상장 체크리스트", "완료", "", "", ds, "22항목 연동"),
        ("D-03", "템플릿", "감리지적_체크리스트 상장/비상장", "완료", "", "", ds, "case_context 상세화"),
        ("D-04", "템플릿", "조서연결_대사_사전 확장", "완료", "", "", ds, "BB·CL·F"),
        ("D-05", "템플릿", "필수절차_카탈로그", "진행중", "", "", "", "KB get_standard_procedures"),
        ("D-06", "템플릿", "검토내역_인정문장", "대기", "", "", "", "qc_review 연동 예정"),
        ("D-07", "템플릿", "골든셋 11건", "완료", "", "", ds, "11/11 통과"),
        ("DEV-01", "개발", "KB 전체 색인·OCR·엑셀", "완료", "", "", ds, "3069문서"),
        ("DEV-02", "개발", "표준절차 KB 연동", "완료", "", "", ds, "knowledge_base·review_engine"),
        ("DEV-03", "개발", "gather_citations·_attach_references", "완료", "", "", ds, "WORKPAPER k=3"),
        ("DEV-04", "개발", "감리지적 enforcement_review", "완료", "", "", ds, ""),
        ("DEV-05", "개발", "골든셋 회귀 스크립트", "완료", "", "", ds, ""),
        ("DEV-06", "개발", "차입금 교차합 BB100/200", "완료", "", "", ds, ""),
        ("DEV-07", "개발", "share_gateway·카톡 랜딩", "완료", "", "", ds, "8506"),
        ("OPS-01", "운영", "Hanul KB 색인 최신화", "완료", "", "", ds, "build_index.py"),
        ("OPS-02", "운영", "사내망 배포 run_server.sh", "완료", "", "", ds, "0.0.0.0:8505"),
        ("OPS-03", "운영", "외부 터널·keep_alive", "완료", "", "", ds, "Cloudflare"),
        ("OPS-04", "운영", "질의회신 빈 PDF 74건", "대기", "", "", "", "원본 교체 필요"),
    ]
    for i, row in enumerate(rows, 2):
        for j, v in enumerate(row, 1):
            ws.cell(i, j, v)
    for i, w in enumerate([8, 12, 36, 10, 10, 12, 12, 28], 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _append_changelog(root: Path, d: date, summary: str) -> None:
    changelog = root / "진행상황" / "변경이력.md"
    line = f"| {d.isoformat()} | (시스템) | {summary} |"
    if changelog.is_file():
        text = changelog.read_text(encoding="utf-8")
        if summary not in text:
            changelog.write_text(text.rstrip() + "\n" + line + "\n", encoding="utf-8")
    else:
        changelog.write_text(
            f"# 변경이력\n\n| 날짜 | 작성자 | 요약 |\n|------|--------|------|\n{line}\n",
            encoding="utf-8",
        )


def deploy_to_hanul_db(d: date | None = None) -> Path:
    """Hanul DB에 개발지침서 폴더·문서 배치."""
    d = d or date.today()
    root = hanul_dev_guide_root()
    progress = root / "진행상황"
    minutes = root / "회의록"
    progress.mkdir(parents=True, exist_ok=True)
    minutes.mkdir(parents=True, exist_ok=True)

    base = guide_basename(d)
    (root / f"{base}.md").write_text(build_development_guide_markdown(d), encoding="utf-8")
    (root / f"{base}.docx").write_bytes(build_development_guide_docx(d))

    readme = f"""# 한울 개발지침서 (협업 허브)

- **최신 개발지침서:** `{base}.docx` / `{base}.md`
- **진행 체크:** `진행상황/개발진행_체크리스트_FY2026.xlsx`
- **변경 기록:** `진행상황/변경이력.md`
- **회의록:** `회의록/` 폴더

## 관련 폴더

- `../자가검토_지침_템플릿/` — QC·운영 xlsx (프로그램 자동 연동)
- `../4대 중점사항 감리대상/` — 금감원·한공회 원문
- `../감사조서 샘플/` — 실무 조서 (KB WORKPAPER)

## 배포·검증

- 로컬: `./run_server.sh` → http://localhost:8505
- 외부: `./run_tunnels.sh` + `./keep_alive.sh`
- KB: `build_index.py` / `golden_set_regression.py`

갱신: {d.isoformat()} — `python dev_guide_export.py`
"""
    (root / "README.md").write_text(readme, encoding="utf-8")

    _append_changelog(
        root,
        d,
        "개발지침서 전면 갱신 — KB학습·감리지적·표준절차·배포·터널·카톡공유 반영",
    )

    (minutes / "README.md").write_text(
        "회의록·결정사항을 `YYYY-MM-DD_주제.md` 형식으로 저장하세요.\n",
        encoding="utf-8",
    )

    xlsx_path = progress / "개발진행_체크리스트_FY2026.xlsx"
    xlsx_path.write_bytes(build_progress_checklist_xlsx(d))

    return root


if __name__ == "__main__":
    p = deploy_to_hanul_db()
    print(p)
