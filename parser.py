"""감사조서 파싱 모듈.

업로드된 조서(PDF·엑셀·워드)에서 텍스트·표를 추출해
규칙엔진이 점검할 수 있는 형태로 정규화합니다.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Any

import pandas as pd


@dataclass
class ParsedDocument:
    """파싱 결과 표준 구조."""

    file_name: str
    file_type: str
    text: str = ""
    tables: list[pd.DataFrame] = field(default_factory=list)
    sheet_names: list[str] = field(default_factory=list)
    page_count: int = 0
    parse_error: str | None = None
    source_files: list[str] = field(default_factory=list)

    @property
    def has_content(self) -> bool:
        return bool(self.text.strip()) or bool(self.tables)

    @property
    def is_merged(self) -> bool:
        return len(self.source_files) > 1


def _ext(file_name: str) -> str:
    return file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""


def _file_stem(file_name: str) -> str:
    return file_name.rsplit(".", 1)[0] if "." in file_name else file_name


def _format_merged_label(file_names: list[str]) -> str:
    if len(file_names) == 1:
        return file_names[0]
    if len(file_names) <= 3:
        return ", ".join(file_names)
    return f"{file_names[0]} 외 {len(file_names) - 1}건"


def merge_documents(docs: list["ParsedDocument"]) -> "ParsedDocument":
    """여러 조서 파싱 결과를 하나의 검토 대상 문서로 합칩니다."""
    valid = [d for d in docs if d.has_content and not d.parse_error]
    if not valid:
        return ParsedDocument(
            file_name="",
            file_type="unknown",
            parse_error="병합할 조서가 없습니다.",
        )
    if len(valid) == 1:
        only = valid[0]
        if not only.source_files:
            only.source_files = [only.file_name]
        return only

    texts: list[str] = []
    tables: list[pd.DataFrame] = []
    sheet_names: list[str] = []
    source_files: list[str] = []
    page_count = 0
    types = {d.file_type for d in valid}

    for doc in valid:
        stem = _file_stem(doc.file_name)
        source_files.append(doc.file_name)
        if doc.text.strip():
            texts.append(f"=== [{doc.file_name}] ===\n{doc.text}")
        page_count += doc.page_count
        for table in doc.tables:
            merged = table.copy()
            orig_source = str(merged.attrs.get("source", "")).strip().strip("/")
            merged.attrs = dict(merged.attrs)
            merged.attrs["source"] = f"{stem}/{orig_source}" if orig_source else stem
            merged.attrs["workpaper_file"] = doc.file_name
            if orig_source and merged.attrs.get("title"):
                merged.attrs["title"] = str(merged.attrs["title"])
            tables.append(merged)
        for name in doc.sheet_names:
            sheet_names.append(f"{stem}/{name}")

    return ParsedDocument(
        file_name=_format_merged_label(source_files),
        file_type=valid[0].file_type if len(types) == 1 else "mixed",
        text="\n\n".join(texts),
        tables=tables,
        sheet_names=sheet_names,
        page_count=page_count,
        source_files=source_files,
    )


def parse_uploads(
    uploads: list[tuple[bytes, str]],
) -> tuple[ParsedDocument | None, list[str]]:
    """여러 업로드 파일을 파싱·병합. (성공 문서, 경고/오류 메시지 목록) 반환."""
    parsed: list[ParsedDocument] = []
    messages: list[str] = []
    for file_bytes, file_name in uploads:
        doc = parse_document(file_bytes, file_name)
        if doc.parse_error:
            messages.append(f"「{file_name}」: {doc.parse_error}")
            continue
        if not doc.has_content:
            messages.append(
                f"「{file_name}」: 텍스트·표를 추출하지 못했습니다. "
                "스캔 PDF이거나 빈 파일일 수 있습니다."
            )
            continue
        doc.source_files = [file_name]
        parsed.append(doc)

    if not parsed:
        return None, messages

    merged = merge_documents(parsed)
    return merged, messages


def parse_document(file_bytes: bytes, file_name: str) -> ParsedDocument:
    """확장자에 따라 적절한 파서로 분기."""
    ext = _ext(file_name)
    try:
        if ext == "pdf":
            return _parse_pdf(file_bytes, file_name)
        if ext in ("xlsx", "xls"):
            return _parse_excel(file_bytes, file_name)
        if ext == "docx":
            return _parse_docx(file_bytes, file_name)
        return ParsedDocument(
            file_name=file_name,
            file_type=ext or "unknown",
            parse_error=f"지원하지 않는 형식입니다: .{ext}",
        )
    except Exception as exc:  # noqa: BLE001 - 사용자에게 원인 표시
        return ParsedDocument(
            file_name=file_name,
            file_type=ext or "unknown",
            parse_error=f"파일을 읽는 중 오류가 발생했습니다: {exc}",
        )


def _parse_pdf(file_bytes: bytes, file_name: str) -> ParsedDocument:
    import pdfplumber

    texts: list[str] = []
    tables: list[pd.DataFrame] = []
    page_count = 0
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        page_count = len(pdf.pages)
        for page_no, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text:
                texts.append(page_text)
            for t_no, raw in enumerate(page.extract_tables() or [], start=1):
                df = _rows_to_df(raw)
                if df is not None:
                    df.attrs["source"] = f"{page_no}페이지 표{t_no}"
                    tables.append(df)
    return ParsedDocument(
        file_name=file_name,
        file_type="pdf",
        text="\n".join(texts),
        tables=tables,
        page_count=page_count,
    )


def _visible_sheet_names(file_bytes: bytes, ext: str) -> set[str] | None:
    """활성(visible) 시트 이름 집합. 판정 불가 시 None(전체 허용)."""
    try:
        if ext == "xlsx":
            import openpyxl

            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True)
            visible = {ws.title for ws in wb.worksheets if ws.sheet_state == "visible"}
            wb.close()
            return visible or None
        import xlrd  # xls

        wb = xlrd.open_workbook(file_contents=file_bytes)
        return {s.name for s in wb.sheets() if s.visibility == 0} or None
    except Exception:  # noqa: BLE001 - 라이브러리 미지원 등은 전체 허용
        return None


def _cell_to_str(value: Any) -> str:
    """셀 값을 텍스트 추출용 문자열로 변환 (숫자·날짜·bool 포함)."""
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except (TypeError, ValueError):
        pass
    if isinstance(value, bool):
        return "TRUE" if value else "FALSE"
    return str(value).strip()


def _safe_join(parts: Any) -> str:
    """join() 전 모든 항목을 문자열로 강제 변환."""
    return " ".join(_cell_to_str(p) for p in parts if _cell_to_str(p))


def _dataframe_to_text(df: pd.DataFrame) -> str:
    """DataFrame 전체를 검색·규칙엔진용 평문으로 직렬화."""
    if df.empty:
        return ""
    lines: list[str] = []
    for _, row in df.iterrows():
        cells = [_cell_to_str(c) for c in row]
        line = " ".join(c for c in cells if c)
        if line:
            lines.append(line)
    return "\n".join(lines)


def _load_workbook_formulas(file_bytes: bytes) -> dict[str, dict[tuple[int, int], str]]:
    """시트별 셀 산식 — {(엑셀행, 엑셀열): '=D9-D10'} (1-based)."""
    try:
        import openpyxl
    except ImportError:
        return {}
    out: dict[str, dict[tuple[int, int], str]] = {}
    try:
        wb = openpyxl.load_workbook(
            io.BytesIO(file_bytes), read_only=True, data_only=False
        )
        for ws in wb.worksheets:
            formulas: dict[tuple[int, int], str] = {}
            for row in ws.iter_rows():
                for cell in row:
                    if cell.data_type == "f" and cell.value:
                        formulas[(cell.row, cell.column)] = str(cell.value)
            if formulas:
                out[ws.title] = formulas
        wb.close()
    except Exception:  # noqa: BLE001
        return {}
    return out


def _parse_excel(file_bytes: bytes, file_name: str) -> ParsedDocument:
    sheet_formulas = _load_workbook_formulas(file_bytes)
    sheets = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None, header=None)

    # 숨겨진 시트는 검토 대상에서 제외 (활성 시트만 리뷰)
    visible = _visible_sheet_names(file_bytes, _ext(file_name))
    if visible is not None:
        sheets = {name: df for name, df in sheets.items() if name in visible}

    tables: list[pd.DataFrame] = []
    text_parts: list[str] = []
    for name, df in sheets.items():
        try:
            text_parts.append(f"[시트: {name}]")
            text_parts.append(_dataframe_to_text(df))
            cleaned = df.dropna(how="all").dropna(axis=1, how="all")
            if not cleaned.empty and cleaned.shape[1] >= 1:
                # reset 전 인덱스가 엑셀 원본 행(0-based) → 엑셀 행번호는 +1
                row_map = [int(i) + 1 for i in cleaned.index]
                col_map = [int(c) + 1 for c in cleaned.columns]  # 엑셀 열번호(1-based)
                cleaned = cleaned.reset_index(drop=True)
                cleaned.attrs["source"] = name
                cleaned.attrs["title"] = _detect_sheet_title(df)
                cleaned.attrs["row_map"] = row_map
                cleaned.attrs["col_map"] = col_map
                cleaned.attrs["cell_formulas"] = sheet_formulas.get(name, {})
                tables.append(cleaned)
        except Exception as exc:  # noqa: BLE001 - 시트 단위 실패는 건너뜀
            text_parts.append(f"[시트: {name}] (일부 추출 실패: {exc})")
    return ParsedDocument(
        file_name=file_name,
        file_type="xlsx",
        text="\n".join(text_parts),
        tables=tables,
        sheet_names=list(sheets.keys()),
    )


_TITLE_CODE_RE = re.compile(r"4000\s*-\s*[A-Z][A-Z0-9.,]*\s+(.+)", re.I)
_TITLE_ACC_RE = re.compile(r"계정과목[^:：]*[:：]\s*(.+)")
_DATE_LIKE_RE = re.compile(r"^\d{4}-\d{2}-\d{2}(?:\s+\d{2}:\d{2}:\d{2})?$")
_TIME_ONLY_RE = re.compile(r"^\d{2}:\d{2}:\d{2}$")


def _detect_sheet_title(df: pd.DataFrame) -> str:
    """시트 상단에서 계정 제목을 추출 (예: '4000-E 재고자산' → '재고자산')."""
    import datetime

    for _, row in df.head(10).iterrows():
        for cell in row:
            if cell is None:
                continue
            if isinstance(cell, (datetime.datetime, datetime.date, datetime.time)):
                continue
            s = str(cell).strip()
            if not s or _DATE_LIKE_RE.match(s):
                continue
            m = _TITLE_ACC_RE.search(s) or _TITLE_CODE_RE.search(s)
            if m:
                title = m.group(1).strip()
                if title and not _TIME_ONLY_RE.match(title):
                    return title
    return ""


def _parse_docx(file_bytes: bytes, file_name: str) -> ParsedDocument:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    tables: list[pd.DataFrame] = []
    for t_no, table in enumerate(doc.tables, start=1):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        df = _rows_to_df(rows)
        if df is not None:
            df.attrs["source"] = f"표{t_no}"
            tables.append(df)
        for row in rows:
            texts.append(" ".join(_cell_to_str(c) for c in row))
    return ParsedDocument(
        file_name=file_name,
        file_type="docx",
        text="\n".join(texts),
        tables=tables,
    )


def _rows_to_df(rows: list[list[Any]] | None) -> pd.DataFrame | None:
    """원시 표 행(2차원 리스트)을 DataFrame으로 변환."""
    if not rows or len(rows) < 2:
        return None
    header = [str(c).strip() if c is not None else "" for c in rows[0]]
    body = [[c if c is not None else "" for c in r] for r in rows[1:]]
    try:
        df = pd.DataFrame(body, columns=header)
    except Exception:  # noqa: BLE001 - 열 수 불일치 등
        df = pd.DataFrame(rows)
    return df if not df.empty else None


_NUM_RE = re.compile(r"-?\d[\d,]*\.?\d*")


def to_number(value: Any) -> float | None:
    """'1,250백만원', '(70)' 같은 셀 값에서 숫자를 추출."""
    if value is None:
        return None
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        num = float(value)
        return None if num != num else num  # NaN 제외
    s = str(value).strip()
    if not s:
        return None
    negative = s.startswith("(") and s.endswith(")")
    m = _NUM_RE.search(s.replace(" ", ""))
    if not m:
        return None
    try:
        num = float(m.group().replace(",", ""))
    except ValueError:
        return None
    return -num if negative else num
