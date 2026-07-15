#!/usr/bin/env python3
"""Streamlit Cloud 고정 URL 배포 교재(Word) 생성."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "Streamlit_고정URL_배포_교재.docx"

# 색상 팔레트
NAVY = RGBColor(0x0F, 0x3E, 0x9C)
GOLD = RGBColor(0xF0, 0xB9, 0x13)
GREEN = RGBColor(0x00, 0x70, 0x47)
RED = RGBColor(0xC0, 0x00, 0x00)
ORANGE = RGBColor(0xED, 0x7D, 0x31)
GRAY = RGBColor(0x55, 0x55, 0x55)
BLUE_BG = "EAF1FC"
GREEN_BG = "E8F5E9"
YELLOW_BG = "FFF8E1"
RED_BG = "FFEBEE"


def _shade_cell(cell, fill_hex: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _run_color(paragraph, text: str, *, color: RGBColor, bold: bool = False, size: int = 11):
    r = paragraph.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return r


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = NAVY if level <= 2 else RGBColor(0x12, 0x23, 0x3F)


def _para(doc: Document, text: str, *, bold: bool = False, color: RGBColor | None = None, size: int = 11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def _box(doc: Document, title: str, body: str, *, fill: str, title_color: RGBColor) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _shade_cell(cell, fill)
    p = cell.paragraphs[0]
    _run_color(p, title + "\n", color=title_color, bold=True, size=12)
    r = p.add_run(body)
    r.font.size = Pt(10.5)
    doc.add_paragraph()


def _steps(doc: Document, steps: list[str]) -> None:
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        _run_color(p, f" ", color=NAVY)
        r = p.add_run(s)
        r.font.size = Pt(11)


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.size = Pt(11)


def build() -> Path:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # 표지
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run_color(t, "한울 감사조서 Smart Reviewer\n", color=NAVY, bold=True, size=26)
    _run_color(t, "Streamlit Cloud 고정 URL 배포 교재\n", color=NAVY, bold=True, size=20)
    _run_color(t, "\n(초보자용 · Mac 꺼도 24시간 접속)\n", color=GOLD, bold=True, size=14)
    _run_color(t, f"\n작성일: {date.today().isoformat()}\n", color=GRAY, size=11)
    _run_color(t, "프로젝트: hanul-002\n", color=GRAY, size=11)
    _run_color(t, "최종 URL: https://hanul-audit-review.streamlit.app\n", color=GREEN, bold=True, size=12)
    doc.add_page_break()

    _heading(doc, "목차", 1)
    toc = [
        "0. 이 교재의 목표 — 왜 이렇게 배포하나?",
        "1. 전체 순서 한눈에 보기 (체크리스트)",
        "2. 사전 준비물",
        "3. GitHub 계정 만들기 (처음부터)",
        "4. GitHub 저장소(Repository) 만들기",
        "5. Mac에 필요한 프로그램 설치",
        "6. 프로젝트 폴더 위치 확인",
        "7. GitHub에 코드 업로드하기",
        "8. Streamlit Cloud에 배포하기",
        "9. 고정 URL 확인 및 교수님께 전달",
        "10. 코드 수정 후 다시 배포하기",
        "11. 자주 발생한 오류와 해결 (FAQ)",
        "부록 A. Mac 임시 URL vs Streamlit 고정 URL",
        "부록 B. 클라우드에서 되는 것 / 안 되는 것",
    ]
    _bullets(doc, toc)
    doc.add_page_break()

    # 0장
    _heading(doc, "0. 이 교재의 목표 — 왜 이렇게 배포하나?", 1)
    _para(doc, "이 교재는 Mac에서 만든 Streamlit 앱을 교수님께 '고정 URL'로 전달하는 방법을 단계별로 설명합니다.")
    _box(
        doc,
        "🎯 최종 목표",
        "교수님이 아래 주소만 클릭하면, 내 Mac이 꺼져 있어도 앱이 작동합니다.\n"
        "https://hanul-audit-review.streamlit.app",
        fill=BLUE_BG,
        title_color=NAVY,
    )
    _box(
        doc,
        "💡 핵심 원리 (비유)",
        "① Mac + 터널 방식 = 집 PC를 켜 두고 CCTV로 보여주기 → Mac 끄면 접속 불가\n"
        "② Streamlit Cloud = 유튜브에 올리고 링크만 공유 → Mac 불필요, URL 고정\n\n"
        "코드는 GitHub에 저장하고, Streamlit이 GitHub에서 코드를 가져가 자기 서버에서 실행합니다.",
        fill=YELLOW_BG,
        title_color=ORANGE,
    )
    doc.add_page_break()

    # 1장
    _heading(doc, "1. 전체 순서 한눈에 보기", 1)
    table = doc.add_table(rows=9, cols=3)
    table.style = "Table Grid"
    headers = ["단계", "무엇을 하나?", "완료"]
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        _shade_cell(c, "0F3E9C")
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    rows = [
        ("1", "GitHub 계정 만들기", "☐"),
        ("2", "GitHub 저장소 hanul-audit-review 생성", "☐"),
        ("3", "Mac에 gh(GitHub CLI) 설치", "☐"),
        ("4", "GitHub_업로드_ship-it.command 실행", "☐"),
        ("5", "share.streamlit.io 에서 Deploy", "☐"),
        ("6", "Python 3.12 + Public 설정", "☐"),
        ("7", "https://hanul-audit-review.streamlit.app 접속 확인", "☐"),
        ("8", "교수님께 URL 전달", "☐"),
    ]
    for ri, (a, b, c) in enumerate(rows, 1):
        table.rows[ri].cells[0].text = a
        table.rows[ri].cells[1].text = b
        table.rows[ri].cells[2].text = c
    doc.add_paragraph()

    # 2장
    _heading(doc, "2. 사전 준비물", 1)
    _bullets(
        doc,
        [
            "Mac 컴퓨터 (배포할 때만 필요, 이후에는 꺼도 됨)",
            "인터넷 연결",
            "이메일 주소 (GitHub·Streamlit 가입용)",
            "프로젝트 폴더: 바탕화면 → ABC05CEO → hanul-002",
            "앱 메인 파일: app.py",
            "KB 데이터: kb_store/hanul_kb.sqlite.gz (약 37MB, 자동 포함)",
        ],
    )
    _box(
        doc,
        "⚠️ 중요: GitHub 계정을 하나만 사용하세요",
        "Streamlit에 로그인한 GitHub 계정과, 코드를 올리는 GitHub 계정이 같아야 합니다.\n"
        "예: github.com/ssscpa77-ship-it 계정으로 통일",
        fill=RED_BG,
        title_color=RED,
    )
    doc.add_page_break()

    # 3장 GitHub 계정
    _heading(doc, "3. GitHub 계정 만들기 (처음부터)", 1)
    _para(doc, "GitHub는 코드를 인터넷에 저장하는 '개발자용 클라우드 저장소'입니다. 무료입니다.")
    _steps(
        doc,
        [
            "브라우저(Safari 또는 Chrome)를 엽니다.",
            "주소창에 https://github.com 입력 후 접속합니다.",
            "오른쪽 상단 [Sign up] (가입) 버튼을 클릭합니다.",
            "이메일 주소를 입력합니다. (예: ssscpa@hanulac.co.kr)",
            "비밀번호를 만듭니다. (8자 이상, 기억하기 쉬운 것)",
            "사용자 이름(User name)을 정합니다. → 이 이름이 주소에 쓰입니다.\n"
            "   예: ssscpa77-ship-it → https://github.com/ssscpa77-ship-it",
            "이메일 인증을 완료합니다. (받은 메일의 Verify 클릭)",
            "로그인 후 우측 상단 프로필 사진 → Settings 에서 이메일이 인증됐는지 확인합니다.",
        ],
    )
    _box(
        doc,
        "✅ 확인 방법",
        "브라우저에서 https://github.com/본인아이디 접속 시 본인 프로필이 보이면 성공.",
        fill=GREEN_BG,
        title_color=GREEN,
    )
    doc.add_page_break()

    # 4장 저장소
    _heading(doc, "4. GitHub 저장소(Repository) 만들기", 1)
    _para(doc, "저장소 = 프로젝트 코드를 담는 폴더. Streamlit이 여기서 코드를 읽습니다.")
    _steps(
        doc,
        [
            "GitHub에 로그인한 상태에서 https://github.com/new 접속",
            "Repository name 에 hanul-audit-review 입력 (정확히 이 이름)",
            "Description 은 비워도 됩니다.",
            "🔘 Private (비공개) 선택 — API 키 등이 노출되지 않도록",
            "❌ Add a README file 체크 해제 (비워 둔 저장소 필요)",
            "❌ Add .gitignore 체크 해제",
            "❌ Choose a license 체크 해제",
            "초록색 [Create repository] 버튼 클릭",
        ],
    )
    _box(
        doc,
        "⚠️ 주의",
        "README를 추가하면 첫 업로드 때 충돌이 날 수 있습니다. 반드시 '빈 저장소'로 만드세요.",
        fill=RED_BG,
        title_color=RED,
    )
    _para(doc, "생성 후 주소 예시:", bold=True)
    _para(doc, "https://github.com/ssscpa77-ship-it/hanul-audit-review", color=GREEN, bold=True)
    doc.add_page_break()

    # 5장 설치
    _heading(doc, "5. Mac에 필요한 프로그램 설치", 1)
    _heading(doc, "5-1. Homebrew (없는 경우)", 2)
    _para(doc, "터미널 앱을 엽니다. (Spotlight에서 '터미널' 검색)")
    _para(doc, "아래 명령을 붙여넣고 Enter:")
    p = doc.add_paragraph()
    r = p.add_run("/bin/bash -c \"$(curl -fsSL https://raw.githubusercontent.com/Homebrew/install/HEAD/install.sh)\"")
    r.font.name = "Menlo"
    r.font.size = Pt(9)

    _heading(doc, "5-2. GitHub CLI (gh) 설치", 2)
    _steps(
        doc,
        [
            "터미널에 brew install gh 입력 후 Enter",
            "설치 완료 후 gh --version 으로 확인",
        ],
    )

    _heading(doc, "5-3. gh 로 GitHub 로그인", 2)
    _steps(
        doc,
        [
            "터미널에 gh auth login 입력",
            "What account? → GitHub.com 선택",
            "Protocol → HTTPS 선택",
            "Authenticate Git with credentials? → Yes",
            "How to authenticate? → Login with a web browser 선택",
            "터미널에 표시되는 one-time code (예: ABCD-1234) 를 복사",
            "Enter 누르면 브라우저가 열림 → https://github.com/login/device",
            "코드 붙여넣기 → Continue → Authorize",
            "⚠️ 반드시 코드를 올릴 GitHub 계정으로 승인 (ssscpa77-ship-it)",
        ],
    )
    _box(
        doc,
        "💡 코드는 어디서 확인?",
        "검은 터미널 창에 'First copy your one-time code: XXXX-XXXX' 라고 표시됩니다.",
        fill=YELLOW_BG,
        title_color=ORANGE,
    )
    doc.add_page_break()

    # 6장 폴더
    _heading(doc, "6. 프로젝트 폴더 위치 확인", 1)
    _para(doc, "Finder에서 아래 경로로 이동합니다:")
    _bullets(
        doc,
        [
            "바탕화면 → ABC05CEO → hanul-002",
            "전체 경로: /Users/본인이름/Desktop/ABC05CEO/hanul-002",
        ],
    )
    _para(doc, "이 폴더에 있어야 할 실행 파일:", bold=True)
    _bullets(
        doc,
        [
            "GitHub_업로드_ship-it.command — GitHub 업로드",
            "고정URL_배포.command — Streamlit 배포 안내",
            "교수님_전달_메시지.txt — 카톡 전달 문구",
            "app.py — 앱 본체",
        ],
    )
    _box(
        doc,
        "더블클릭이 안 될 때",
        "파일 우클릭 → [열기] → [열기] 다시 클릭 (보안 경고 우회)",
        fill=YELLOW_BG,
        title_color=ORANGE,
    )
    doc.add_page_break()

    # 7장 업로드
    _heading(doc, "7. GitHub에 코드 업로드하기", 1)
    _para(doc, "4장에서 저장소를 만든 뒤, Mac의 코드를 GitHub로 보냅니다.")
    _steps(
        doc,
        [
            "Finder에서 hanul-002 폴더를 엽니다.",
            "GitHub_업로드_ship-it.command 파일을 더블클릭합니다.",
            "검은 터미널 창이 열리면 '저장소 만든 뒤 엔터' 안내가 나옵니다.",
            "4장에서 저장소를 이미 만들었다면 Enter 키를 누릅니다.",
            "gh auth login 이 실행되면 5-3장과 같이 로그인합니다.",
            "업로드가 1~3분 진행됩니다. '✅ 업로드 완료!' 가 보이면 성공.",
            "브라우저에서 https://github.com/본인아이디/hanul-audit-review 접속",
            "app.py, requirements.txt, kb_store/hanul_kb.sqlite.gz 파일이 보이는지 확인",
        ],
    )
    _box(
        doc,
        "✅ 업로드 성공 확인",
        "GitHub 저장소 페이지에 app.py와 hanul_kb.sqlite.gz(약 37MB)가 보이면 OK.",
        fill=GREEN_BG,
        title_color=GREEN,
    )
    _box(
        doc,
        "❌ 실패 시",
        "• '저장소 없음' → 4장 저장소 생성 여부 확인\n"
        "• '권한 없음' → gh auth login 을 올바른 계정으로 다시 실행\n"
        "• 다른 계정에 올라감 → Streamlit 로그인 계정과 GitHub 계정 통일",
        fill=RED_BG,
        title_color=RED,
    )
    doc.add_page_break()

    # 8장 Streamlit
    _heading(doc, "8. Streamlit Cloud에 배포하기", 1)
    _para(doc, "Streamlit Cloud는 무료로 Streamlit 앱을 24시간 호스팅해 줍니다.")
    _heading(doc, "8-1. Streamlit 가입 및 GitHub 연결", 2)
    _steps(
        doc,
        [
            "https://share.streamlit.io 접속",
            "Continue with GitHub 클릭",
            "⚠️ 7장과 동일한 GitHub 계정으로 로그인",
            "Authorize streamlit 권한 승인",
        ],
    )

    _heading(doc, "8-2. 새 앱 배포 (Deploy)", 2)
    _steps(
        doc,
        [
            "[Create app] 또는 https://share.streamlit.io/deploy 접속",
            "Repository: 본인아이디/hanul-audit-review 선택",
            "Branch: main",
            "Main file path: app.py",
            "App URL (서브도메인): hanul-audit-review",
            "⚙️ Advanced settings 클릭",
            "   → Python version: 3.12 선택 (3.13 아님!)",
            "   → Secrets: 비워도 됨 (AI는 나중에 추가 가능)",
            "Deploy 버튼 클릭",
            "5~10분 대기 (Installing requirements → Running)",
        ],
    )

    _heading(doc, "8-3. 배포 후 필수 설정", 2)
    _steps(
        doc,
        [
            "My apps → hanul-audit-review 클릭",
            "Settings (또는 Manage app → Settings)",
            "App visibility → Public (공개) 선택 ← 필수!",
            "Save 후 Reboot app",
        ],
    )
    _box(
        doc,
        "⚠️ Public 이 아니면",
        "'You do not have access to this app' 오류가 납니다. 교수님도 접속 불가.",
        fill=RED_BG,
        title_color=RED,
    )

    _heading(doc, "8-4. 고정URL_배포.command 로 안내 받기", 2)
    _para(doc, "hanul-002 폴더의 고정URL_배포.command 를 더블클릭하면 배포 페이지가 자동으로 열리고, 클릭 순서가 안내됩니다.")
    doc.add_page_break()

    # 9장 확인
    _heading(doc, "9. 고정 URL 확인 및 교수님께 전달", 1)
    _steps(
        doc,
        [
            "브라우저에서 https://hanul-audit-review.streamlit.app 접속",
            "'한울회계법인 · 감사조서 자가검토' 화면이 보이면 성공",
            "Mac을 끄거나 잠자기 모드로 전환",
            "휴대폰( LTE/5G )에서 같은 URL 접속 → 되면 완료",
            "hanul-002/교수님_전달_메시지.txt 내용을 카톡으로 전달",
        ],
    )
    _box(
        doc,
        "📱 교수님 전달 메시지 (복사용)",
        "[ABC 5기 신성섭 · 감사조서 Smart Reviewer]\n\n"
        "한울회계법인 감사조서 자가검토 시스템 (검증용)\n\n"
        "▶ 접속 링크 (고정)\n"
        "https://hanul-audit-review.streamlit.app\n\n"
        "• PDF·엑셀 조서 업로드 → 리뷰노트 자동 생성\n"
        "• 규칙엔진 + Hanul DB + AI 심층 분석\n"
        "• 절차누락·감리지적·4대중점 자동 점검\n\n"
        "※ PC·태블릿 브라우저(Chrome, Safari) 권장\n"
        "※ 검증용 샘플 조서로 테스트 부탁드립니다.",
        fill=BLUE_BG,
        title_color=NAVY,
    )
    doc.add_page_break()

    # 10장 재배포
    _heading(doc, "10. 코드 수정 후 다시 배포하기", 1)
    _para(doc, "앱 코드를 수정한 뒤 GitHub에 올리면 Streamlit이 자동으로 다시 배포합니다. URL은 그대로입니다.")
    _steps(
        doc,
        [
            "Mac에서 코드 수정",
            "GitHub_업로드_ship-it.command 더블클릭",
            "업로드 완료 후 5~10분 대기",
            "Streamlit My apps 에서 Reboot app (필요 시)",
            "같은 URL로 접속 확인",
        ],
    )
    doc.add_page_break()

    # 11장 FAQ
    _heading(doc, "11. 자주 발생한 오류와 해결 (FAQ)", 1)
    faqs = [
        (
            "Error installing requirements",
            "원인: Python 3.13 + 패키지 호환 문제\n"
            "해결: Advanced settings → Python 3.12 선택 후 앱 삭제·재배포",
        ),
        (
            "Error running app",
            "원인: .streamlit/config.toml 에 port=8505 설정 (클라우드 충돌)\n"
            "해결: config.toml 에서 [server] 섹션 제거 후 GitHub 재업로드",
        ),
        (
            "You do not have access to this app",
            "원인: 앱이 Private 이거나, 삭제·재배포 후 URL 꼬임\n"
            "해결: Settings → Public 설정 / 또는 다른 App URL로 먼저 배포 후 변경",
        ),
        (
            "Repository does not exist (Streamlit)",
            "원인: Streamlit 로그인 계정 ≠ 코드가 있는 GitHub 계정\n"
            "해결: 동일 계정으로 통일 후 재배포",
        ),
        (
            "GitHub one-time code 어디?",
            "터미널(검은 창)에 XXXX-XXXX 표시 → github.com/login/device 에 입력",
        ),
        (
            "보조 URL은 되는데 hanul-audit-review 는 안 됨",
            "삭제·재배포 후 생기는 Streamlit 버그. Settings에서 App URL을 hanul-audit-review 로 변경",
        ),
        (
            "첫 접속이 느림",
            "무료 플랜은 잠자기 상태. 첫 접속 시 30초~2분 깨어남. KB 복원에 1~2분 추가 가능",
        ),
    ]
    for title, body in faqs:
        _box(doc, f"❓ {title}", body, fill=YELLOW_BG, title_color=ORANGE)

    doc.add_page_break()

    # 부록 A
    _heading(doc, "부록 A. Mac 임시 URL vs Streamlit 고정 URL", 1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    data = [
        ("비교", "Mac + 터널", "Streamlit Cloud"),
        ("비용", "0원", "0원"),
        ("Mac 필요", "항상 켜 둬야 함", "불필요"),
        ("URL", "매번 바뀜", "고정"),
        ("교수님 전달", "비추천", "✅ 추천"),
    ]
    for ri, row in enumerate(data):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = val
            if ri == 0:
                _shade_cell(t.rows[ri].cells[ci], "0F3E9C")

    doc.add_paragraph()
    _heading(doc, "부록 B. 클라우드에서 되는 것 / 안 되는 것", 1)
    t2 = doc.add_table(rows=5, cols=2)
    t2.style = "Table Grid"
    bdata = [
        ("✅ 클라우드에서 됨", "❌ Mac에만 있음"),
        ("앱 접속·조서 업로드", "Hanul DB-SSS PDF 원문"),
        ("리뷰노트·엑셀 다운로드", "로컬 폴더 경로 파일"),
        ("KB 검색 (gzip DB)", ""),
        ("규칙엔진 점검", ""),
    ]
    for ri, row in enumerate(bdata):
        for ci, val in enumerate(row):
            t2.rows[ri].cells[ci].text = val
            if ri == 0:
                _shade_cell(t2.rows[ri].cells[ci], "EAF1FC" if ci == 0 else "FFEBEE")

    doc.add_paragraph()
    _para(doc, "— 끝 —", bold=True, color=NAVY, size=14)
    _para(doc, "문의: hanul-002 프로젝트 폴더의 deploy/STREAMLIT_CLOUD.md 참고", color=GRAY)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(path)
